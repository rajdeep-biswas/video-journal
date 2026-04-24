#!/usr/bin/env python3
"""
Pipeline: MOV files -> Bengali SRT + transcriptions + multi-system AI journal Word doc
  - Whisper transcribes in Bengali (Unicode)
  - Anthropic + OpenAI independently analyze; Anthropic synthesizes the best result
  - Journal: Calibri font, emojis, transcript excerpts per section, smart screenshots
"""

import os
import ssl
import json
import time
import tempfile
import subprocess
from io import BytesIO
from pathlib import Path

ssl._create_default_https_context = ssl._create_unverified_context  # corporate proxy fix

import whisper
from PIL import Image
from tqdm import tqdm
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── Config ───────────────────────────────────────────────────────────────────
# Transcription: "gemini" (recommended — handles Bengali/English code-switching)
#                "whisper" (local, offline fallback)
TRANSCRIPTION_PROVIDER = "gemini"
GEMINI_TRANSCRIBE_MODEL = "gemini-2.0-flash"

WHISPER_MODEL    = "small"       # only used if TRANSCRIPTION_PROVIDER = "whisper"
WHISPER_LANGUAGE = "bn"          # only used if TRANSCRIPTION_PROVIDER = "whisper"

CHUNK_DURATION   = 300           # 5 minutes per chunk (seconds)
CHUNK_OVERLAP    = 10            # overlap between consecutive chunks (seconds)
MAX_TOKENS       = 8192          # plenty of room for 5-min word-for-word transcripts

# ─── Paths ────────────────────────────────────────────────────────────────────
BASE               = Path(__file__).parent
SOURCE_DIR         = BASE / "source"
SUBTITLES_DIR      = BASE / "subtitles"
TRANSCRIPTIONS_DIR = BASE / "transcriptions"
JOURNAL_DIR        = BASE / "journal"
SCREENSHOTS_DIR    = BASE / "screenshots"

for d in [SUBTITLES_DIR, TRANSCRIPTIONS_DIR, JOURNAL_DIR, SCREENSHOTS_DIR]:
    d.mkdir(exist_ok=True)

# ─── Key loader ───────────────────────────────────────────────────────────────
def load_key(provider: str) -> str:
    env_map = {"openai": "OPENAI_API_KEY", "anthropic": "ANTHROPIC_API_KEY"}
    key = os.environ.get(env_map.get(provider, ""), "").strip()
    if not key:
        kf = BASE / "keys.json"
        if kf.exists():
            key = json.loads(kf.read_text()).get(provider, "").strip()
    return key

# ─── SRT ──────────────────────────────────────────────────────────────────────
def fmt_srt_ts(s: float) -> str:
    ms = int((s % 1) * 1000)
    s  = int(s)
    return f"{s//3600:02d}:{(s%3600)//60:02d}:{s%60:02d},{ms:03d}"

def write_srt(segments: list, path: Path):
    lines = [
        f"{i}\n{fmt_srt_ts(seg['start'])} --> {fmt_srt_ts(seg['end'])}\n{seg['text'].strip()}\n"
        for i, seg in enumerate(segments, 1)
    ]
    path.write_text("\n".join(lines), encoding="utf-8")

# ─── Video helpers ────────────────────────────────────────────────────────────
def get_duration(mov: Path) -> float:
    r = subprocess.run(
        ["ffprobe", "-v", "error", "-show_entries", "format=duration",
         "-of", "default=noprint_wrappers=1", str(mov)],
        capture_output=True, text=True,
    )
    for line in r.stdout.splitlines():
        if "duration=" in line:
            return float(line.split("=")[1])
    return 0.0

def extract_frame(mov: Path, ts: float, out: Path) -> bool:
    r = subprocess.run(
        ["ffmpeg", "-y", "-ss", str(ts), "-i", str(mov),
         "-vframes", "1", "-q:v", "3", str(out)],
        capture_output=True,
    )
    return r.returncode == 0 and out.exists()

def _chunk_spans(duration: float):
    """Yield (abs_start, audio_length, seam_left, seam_right) for each chunk.

    seam_left / seam_right are the absolute-time window we keep from this chunk,
    placed at the midpoint of each overlap so there are no gaps or duplicates.
    """
    i = 0
    while True:
        abs_start   = i * CHUNK_DURATION
        if abs_start >= duration:
            break
        audio_len   = min(CHUNK_DURATION + CHUNK_OVERLAP, duration - abs_start)
        seam_left   = abs_start + (CHUNK_OVERLAP / 2 if i > 0 else 0)
        next_start  = (i + 1) * CHUNK_DURATION
        seam_right  = next_start + CHUNK_OVERLAP / 2 if next_start < duration else float("inf")
        yield abs_start, audio_len, seam_left, seam_right
        i += 1

GEMINI_TRANSCRIBE_PROMPT = """\
Transcribe this audio with high accuracy.

The speaker freely mixes Bengali and English (code-switching is normal — do not force one language).
- Bengali speech → Bengali Unicode script (বাংলা), NOT romanised/transliterated
- English speech → English exactly as spoken
- Filler sounds (উম, এই, hmm, uh) → include them, they're part of the journal

Return ONLY a JSON array (no markdown fences), each element:
{"start": <seconds from start of THIS clip>, "end": <seconds>, "text": "..."}

Segments should be 3–15 seconds each. Cover every word — do not skip anything."""

def _gemini_transcribe_chunk(gemini_client, chunk_wav: Path, abs_start: float,
                              seam_left: float, seam_right: float) -> list:
    """Upload one audio chunk to Gemini, get segments, adjust to absolute timestamps."""
    from google.genai import types as gt

    uploaded = gemini_client.files.upload(file=chunk_wav)
    try:
        resp = gemini_client.models.generate_content(
            model=GEMINI_TRANSCRIBE_MODEL,
            contents=[
                gt.Part.from_uri(file_uri=uploaded.uri, mime_type="audio/wav"),
                GEMINI_TRANSCRIBE_PROMPT,
            ],
            config=gt.GenerateContentConfig(max_output_tokens=MAX_TOKENS),
        )
    finally:
        try:
            gemini_client.files.delete(name=uploaded.name)
        except Exception:
            pass

    raw = resp.text
    start_idx = raw.find("[")
    end_idx   = raw.rfind("]") + 1
    segs = json.loads(raw[start_idx:end_idx])

    kept = []
    for seg in segs:
        seg["start"] += abs_start
        seg["end"]   += abs_start
        if seam_left <= seg["start"] < seam_right:
            kept.append(seg)
    return kept

def _whisper_transcribe_chunk(wmodel, chunk_wav: Path, abs_start: float,
                               seam_left: float, seam_right: float,
                               language: str = None) -> list:
    """Transcribe one audio chunk with Whisper, adjust to absolute timestamps."""
    kwargs = {"verbose": False, "task": "transcribe"}
    if language:
        kwargs["language"] = language
    result = wmodel.transcribe(str(chunk_wav), **kwargs)
    kept = []
    for seg in result["segments"]:
        seg = dict(seg)
        seg["start"] += abs_start
        seg["end"]   += abs_start
        if seam_left <= seg["start"] < seam_right:
            kept.append(seg)
    return kept

def transcribe_chunked(mov: Path, duration: float,
                        gemini_client=None, wmodel=None) -> tuple:
    """Transcribe in CHUNK_DURATION chunks with CHUNK_OVERLAP-second overlaps.

    Uses Gemini if gemini_client is provided, otherwise falls back to wmodel (Whisper).
    Returns (all_segments, full_text, chunk_transcripts).
    Temp WAV files live only inside a TemporaryDirectory — nothing written to disk permanently.
    """
    spans  = list(_chunk_spans(duration))
    all_segs: list = []
    chunk_transcripts: list = []
    provider_label = "Gemini" if gemini_client else "Whisper"

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        for idx, (abs_start, audio_len, seam_left, seam_right) in enumerate(
            tqdm(spans, desc=f"  {provider_label} chunks", unit="chunk")
        ):
            chunk_wav = tmp / f"chunk_{idx:03d}.wav"
            subprocess.run(
                ["ffmpeg", "-y", "-ss", str(abs_start), "-t", str(audio_len),
                 "-i", str(mov), "-ac", "1", "-ar", "16000", "-f", "wav", str(chunk_wav)],
                capture_output=True, check=True,
            )

            if gemini_client:
                kept = _gemini_transcribe_chunk(
                    gemini_client, chunk_wav, abs_start, seam_left, seam_right
                )
            else:
                kept = _whisper_transcribe_chunk(
                    wmodel, chunk_wav, abs_start, seam_left, seam_right,
                    language=WHISPER_LANGUAGE,
                )

            all_segs.extend(kept)
            chunk_text = " ".join(s["text"].strip() for s in kept)
            chunk_transcripts.append((abs_start, audio_len, chunk_text))

    full_text = " ".join(s["text"].strip() for s in all_segs)
    return all_segs, full_text, chunk_transcripts

def segments_for_range(segments: list, t_start: float, t_end: float) -> str:
    """Return concatenated transcript text covering [t_start, t_end]."""
    chunks = [
        seg["text"].strip()
        for seg in segments
        if seg["end"] >= t_start - 2 and seg["start"] <= t_end + 2
    ]
    return " ".join(chunks)

# ─── Analysis schema ──────────────────────────────────────────────────────────
SCHEMA = """\
Return ONLY a JSON object (no markdown fences) with this exact structure:
{
  "title": "short evocative title for this recording session",
  "summary": "2–3 rich sentences capturing the essence of the session",
  "mood_and_energy": "vivid description of the vibe, mental state, and emotional texture",
  "topics": [
    {
      "heading": "section heading",
      "description": "DETAILED description — for introspective sections write 3–5 sentences capturing the emotional depth, the specific thoughts, doubts, realisations; for activity sections describe exactly what was done/shown step by step",
      "is_introspective": true or false,
      "screenshot_worthy": true or false — true ONLY if something visually interesting is on screen (not just a talking face),
      "timestamp_start": <seconds>,
      "timestamp_end": <seconds>
    }
  ],
  "key_moments": [
    {
      "timestamp": <seconds>,
      "description": "what is happening — pick 2–4 moments max, only where something is visually interesting or physically shown",
      "screenshot_worthy": true
    }
  ],
  "introspections": [
    "Full, rich sentence capturing a deep thought or realisation — include the emotional context and what triggered it"
  ],
  "todos": [
    { "item": "concrete task they want to do", "context": "why or when" }
  ],
  "loose_threads": [
    { "thought": "unfinished thought or open question", "context": "what triggered it" }
  ]
}"""

SYNTHESIS_SCHEMA = """\
You are given two independent AI analyses (A and B) of the same video journal transcript.
Synthesise the BEST possible version:
- For descriptions: pick whichever is richer and more insightful; merge if both add value
- For todos and loose_threads: take the UNION — include items from either analysis
- For introspections: take the union and deduplicate, keeping the most detailed phrasing
- For key_moments: keep only 2–4 total, only screenshot_worthy=true ones
- For mood/title/summary: pick whichever captures the session better

Return ONLY a JSON object matching the exact same schema as the inputs (no markdown fences)."""

def build_chunk_prompt(video_name: str, chunk_start: float, chunk_end: float,
                        total_duration: float, transcript: str) -> str:
    def fmt(s): m, sec = divmod(int(s), 60); return f"{m}:{sec:02d}"
    return f"""You are analyzing a 5-minute segment of a personal video journal.
The person recorded themselves — talking, thinking aloud, showing their screen, working, introspecting.

Video: "{video_name}"
Segment: {fmt(chunk_start)} – {fmt(chunk_end)}  (of {fmt(total_duration)} total)

Transcript for this segment:
---
{transcript}
---

Instructions:
- All timestamps in your JSON must be ABSOLUTE (seconds from start of full video, not this segment).
  So add {int(chunk_start)} to any relative timestamps.
- Introspective sections: write 3–5 sentences capturing emotional depth, specific thoughts, doubts, realisations.
- Activity sections: describe exactly what was done/shown step by step.
- For loose_threads: mid-sentence stops, unanswered questions, ideas started but not resolved.
- screenshot_worthy=true ONLY when something is physically shown / screen content changes — NOT talking-face moments.
- key_moments: max 2 per segment, all must be screenshot_worthy=true.

{SCHEMA}"""

def build_synthesis_prompt(video_name: str, total_duration: float,
                            chunk_analyses: list) -> str:
    """Prompt to merge N chunk analyses into one coherent video-level analysis."""
    def fmt(s): m, sec = divmod(int(s), 60); return f"{m}:{sec:02d}"
    chunks_json = "\n\n".join(
        f"--- Segment {i+1} ({fmt(start)}–{fmt(start+dur)}) ---\n"
        f"Anthropic: {json.dumps(a, ensure_ascii=False)}\n"
        f"OpenAI:    {json.dumps(b, ensure_ascii=False)}"
        for i, (start, dur, a, b) in enumerate(chunk_analyses)
    )
    return f"""You are given dual AI analyses (Anthropic + OpenAI) for each 5-minute segment of a
personal video journal recording. Synthesise everything into ONE coherent, complete journal entry.

Video: "{video_name}"  |  Total duration: {fmt(total_duration)}

Per-segment analyses:
{chunks_json}

Synthesis rules:
- topics: merge into a flowing narrative of the full session; preserve all introspective depth
- todos + loose_threads: UNION from all segments and both models — miss nothing
- introspections: union, deduplicate, keep the richest phrasing
- key_moments: max 4 total across the whole video, only screenshot_worthy=true
- title/summary/mood: synthesise across the full arc of the session

{SCHEMA}"""

# ─── AI helpers ───────────────────────────────────────────────────────────────
def _parse(raw: str) -> dict:
    return json.loads(raw[raw.find("{"):raw.rfind("}")+1])

def analyze_anthropic_chunk(client, video_name, chunk_start, chunk_end,
                             total_duration, transcript) -> dict:
    prompt = build_chunk_prompt(video_name, chunk_start, chunk_end, total_duration, transcript)
    resp = client.messages.create(
        model="claude-opus-4-7",
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": prompt}],
    )
    return _parse(resp.content[0].text)

def analyze_openai_chunk(client, video_name, chunk_start, chunk_end,
                          total_duration, transcript) -> dict:
    prompt = build_chunk_prompt(video_name, chunk_start, chunk_end, total_duration, transcript)
    resp = client.chat.completions.create(
        model="gpt-4o",
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": prompt}],
    )
    return _parse(resp.choices[0].message.content)

def synthesize_all_chunks(anthropic_client, video_name, total_duration,
                           chunk_analyses: list) -> dict:
    """chunk_analyses: list of (chunk_start, chunk_dur, ant_analysis, oai_analysis)."""
    prompt = build_synthesis_prompt(video_name, total_duration, chunk_analyses)
    resp = anthropic_client.messages.create(
        model="claude-opus-4-7",
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": prompt}],
    )
    return _parse(resp.content[0].text)

# ─── Word doc helpers ─────────────────────────────────────────────────────────
FONT_NAME   = "Calibri"
COLOR_GRAY  = RGBColor(0x77, 0x77, 0x77)
COLOR_BLUE  = RGBColor(0x1F, 0x5C, 0x99)
COLOR_PURPLE= RGBColor(0x66, 0x44, 0x99)
COLOR_TEAL  = RGBColor(0x1A, 0x7A, 0x72)
COLOR_AMBER = RGBColor(0xB8, 0x76, 0x00)

def set_doc_fonts(doc: Document):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3",
                        "List Bullet", "List Bullet 2", "Quote"]:
        try:
            style = doc.styles[style_name]
            style.font.name = FONT_NAME
            rPr = style.element.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), FONT_NAME)
            rFonts.set(qn("w:hAnsi"), FONT_NAME)
            rFonts.set(qn("w:cs"), FONT_NAME)
            rPr.insert(0, rFonts)
        except Exception:
            pass

def add_run(para, text: str, bold=False, italic=False, color=None, size=None):
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.name = FONT_NAME
    if color: r.font.color.rgb = color
    if size:  r.font.size = Pt(size)
    return r

def add_heading(doc, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
    return h

def insert_screenshot(doc: Document, img_path: Path, caption: str):
    """Insert image via Pillow (fixes python-docx JPEG detection bug)."""
    try:
        buf = BytesIO()
        Image.open(img_path).convert("RGB").save(buf, format="JPEG", quality=85)
        buf.seek(0)
        doc.add_picture(buf, width=Inches(5.2))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name  = FONT_NAME
            r.font.size  = Pt(8)
            r.font.color.rgb = COLOR_GRAY
    except Exception as e:
        print(f"    [screenshot error: {e}]")

def add_transcript_block(doc: Document, text: str):
    """Add a styled transcript excerpt block."""
    if not text.strip():
        return
    p = doc.add_paragraph(style="Quote") if "Quote" in [s.name for s in doc.styles] else doc.add_paragraph()
    run = p.add_run(f'"{text.strip()}"')
    run.font.name   = FONT_NAME
    run.font.size   = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = COLOR_GRAY
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)

def ts_label(s: float) -> str:
    m, sec = divmod(int(s), 60)
    return f"{m}:{sec:02d}"

# ─── Journal builder ──────────────────────────────────────────────────────────
def build_journal(entries: list, output_path: Path):
    doc = Document()
    set_doc_fonts(doc)

    # Cover
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title_p, "📓  Personal Video Journal", bold=True, size=22, color=COLOR_BLUE)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(date_p, "April 26, 2023", color=COLOR_GRAY, size=11)
    doc.add_paragraph()

    for entry in entries:
        name     = entry["video_name"]
        duration = entry["duration"]
        analysis = entry["analysis"]
        shots    = entry["screenshots"]        # list of (ts, Path)
        segments = entry["segments"]           # whisper segments for transcript excerpts

        used_shots: set = set()

        def best_shot(ts: float, require_worthy: bool = True) -> Path | None:
            candidates = [(t, p) for t, p in shots if p.exists()]
            if not candidates:
                return None
            _, path = min(candidates, key=lambda x: abs(x[0] - ts))
            if str(path) in used_shots:
                return None
            return path

        def use_shot(path: Path):
            used_shots.add(str(path))

        # ── Session heading ────────────────────────────────────────────────────
        add_heading(doc, f"🎬  {analysis.get('title', name)}", level=1)

        meta = doc.add_paragraph()
        add_run(meta, f"{name}  ·  {int(duration//60)}m {int(duration%60)}s",
                italic=True, color=COLOR_GRAY, size=9)

        if analysis.get("mood_and_energy"):
            vibe = doc.add_paragraph()
            add_run(vibe, f"✨  {analysis['mood_and_energy']}", italic=True, color=COLOR_TEAL)

        doc.add_paragraph()

        # ── Overview ───────────────────────────────────────────────────────────
        add_heading(doc, "📋  Overview", level=2)
        summary_p = doc.add_paragraph()
        add_run(summary_p, analysis.get("summary", ""))

        # ── Topics ────────────────────────────────────────────────────────────
        topics = analysis.get("topics", [])
        if topics:
            add_heading(doc, "🗂️  What I Was Doing", level=2)
            for topic in topics:
                is_intro = topic.get("is_introspective", False)
                icon     = "💭" if is_intro else "⚡"
                add_heading(doc, f"{icon}  {topic['heading']}", level=3)

                desc_p = doc.add_paragraph()
                add_run(desc_p, topic["description"])

                # Transcript excerpt for this section
                t0 = topic.get("timestamp_start", 0)
                t1 = topic.get("timestamp_end", t0 + 60)
                excerpt = segments_for_range(segments, t0, t1)
                add_transcript_block(doc, excerpt)

                # Screenshot — only if AI marked it worthy
                if topic.get("screenshot_worthy"):
                    shot = best_shot(t0)
                    if shot:
                        use_shot(shot)
                        insert_screenshot(doc, shot, f"[{ts_label(t0)}]  {topic['heading']}")

        # ── Introspections ─────────────────────────────────────────────────────
        introspections = analysis.get("introspections", [])
        if introspections:
            add_heading(doc, "🪞  Reflections & Introspections", level=2)
            for thought in introspections:
                p = doc.add_paragraph(style="List Bullet")
                add_run(p, thought, italic=True, color=COLOR_PURPLE)

        # ── Key moments ────────────────────────────────────────────────────────
        key_moments = [km for km in analysis.get("key_moments", [])
                       if km.get("screenshot_worthy", True)]
        if key_moments:
            add_heading(doc, "📸  Key Moments", level=2)
            for km in key_moments:
                ts   = km.get("timestamp", 0)
                desc = km.get("description", "")
                p = doc.add_paragraph(style="List Bullet")
                add_run(p, f"[{ts_label(ts)}]  {desc}", color=COLOR_BLUE)

                shot = best_shot(ts)
                if shot:
                    use_shot(shot)
                    insert_screenshot(doc, shot, f"[{ts_label(ts)}]  {desc}")

        # ── TODOs ─────────────────────────────────────────────────────────────
        todos = analysis.get("todos", [])
        if todos:
            add_heading(doc, "✅  TODO", level=2)
            note_p = doc.add_paragraph()
            add_run(note_p, "Action items from this session:", bold=True)
            for todo in todos:
                p = doc.add_paragraph(style="List Bullet")
                add_run(p, f"☐  {todo['item']}", bold=True)
                if todo.get("context"):
                    ctx = doc.add_paragraph()
                    add_run(ctx, f"    ↳ {todo['context']}", color=COLOR_GRAY, size=9.5)

        # ── Loose Threads ─────────────────────────────────────────────────────
        threads = analysis.get("loose_threads", [])
        if threads:
            add_heading(doc, "🧵  Loose Threads", level=2)
            note_p = doc.add_paragraph()
            add_run(note_p, "Thoughts left unfinished — return to these:",
                    italic=True, color=COLOR_PURPLE)
            for thread in threads:
                p = doc.add_paragraph(style="List Bullet")
                add_run(p, f"◌  {thread['thought']}", color=COLOR_AMBER)
                if thread.get("context"):
                    ctx = doc.add_paragraph()
                    add_run(ctx, f"    ↳ {thread['context']}", color=COLOR_GRAY, size=9.5)

        doc.add_page_break()

    doc.save(str(output_path))
    print(f"  Saved: {output_path}")

# ─── Logging helpers ──────────────────────────────────────────────────────────
def step(msg: str):
    print(f"\n  ▶  {msg}", flush=True)

def done(msg: str, elapsed: float):
    print(f"  ✓  {msg}  ({elapsed:.1f}s)", flush=True)

# ─── Main ─────────────────────────────────────────────────────────────────────
def main():
    # Load API keys
    anthropic_key = load_key("anthropic")
    openai_key    = load_key("openai")

    if not anthropic_key:
        print("ERROR: anthropic key not found in keys.json or ANTHROPIC_API_KEY env var")
        return
    if not openai_key:
        print("ERROR: openai key not found in keys.json or OPENAI_API_KEY env var")
        return

    import anthropic as anthropic_sdk
    from openai import OpenAI

    ant_client = anthropic_sdk.Anthropic(api_key=anthropic_key)
    oai_client = OpenAI(api_key=openai_key)

    # ── Transcription client ───────────────────────────────────────────────────
    gemini_client = None
    wmodel        = None
    t0_total = time.time()

    if TRANSCRIPTION_PROVIDER == "gemini":
        gemini_key = load_key("gemini")
        if not gemini_key:
            print("ERROR: gemini key not found in keys.json — needed for transcription")
            return
        from google import genai as google_genai
        gemini_client = google_genai.Client(api_key=gemini_key)
        print(f"\nTranscription: Gemini ({GEMINI_TRANSCRIBE_MODEL})", flush=True)
    else:
        print(f"\nLoading Whisper model ({WHISPER_MODEL})...", flush=True)
        t0 = time.time()
        import whisper as whisper_mod
        wmodel = whisper_mod.load_model(WHISPER_MODEL)
        print(f"  Model loaded  ({time.time()-t0:.1f}s)", flush=True)

    mov_files = sorted(
        list(SOURCE_DIR.glob("*.mov")) + list(SOURCE_DIR.glob("*.MOV")),
        key=lambda p: p.name,
    )

    if not mov_files:
        print("No .mov files found in source/")
        return

    print(f"\nFound {len(mov_files)} video(s) to process.", flush=True)
    entries = []

    for idx, mov in enumerate(mov_files, 1):
        print(f"\n{'━'*60}", flush=True)
        print(f"[{idx}/{len(mov_files)}]  {mov.name}", flush=True)
        print(f"{'━'*60}", flush=True)

        stem     = mov.stem
        duration = get_duration(mov)
        print(f"  Duration: {int(duration//60)}m {int(duration%60)}s", flush=True)

        # ── Chunked transcription ─────────────────────────────────────────────
        n_chunks = max(1, int(duration // CHUNK_DURATION) + (1 if duration % CHUNK_DURATION else 0))
        provider_label = "Gemini" if gemini_client else f"Whisper ({WHISPER_LANGUAGE})"
        step(f"{provider_label} transcription — {n_chunks} chunk(s) × 5 min")
        t0 = time.time()
        segments, full_text, chunk_transcripts = transcribe_chunked(
            mov, duration, gemini_client=gemini_client, wmodel=wmodel
        )
        done(f"Transcribed {n_chunks} chunk(s)", time.time() - t0)

        if not full_text:
            print("  ⚠  No speech detected — skipping.", flush=True)
            (TRANSCRIPTIONS_DIR / f"{stem}.txt").write_text("(no speech detected)", encoding="utf-8")
            (SUBTITLES_DIR / f"{stem}.srt").write_text("", encoding="utf-8")
            continue

        # ── Save transcript & SRT ─────────────────────────────────────────────
        (TRANSCRIPTIONS_DIR / f"{stem}.txt").write_text(full_text, encoding="utf-8")
        write_srt(segments, SUBTITLES_DIR / f"{stem}.srt")
        print(f"  💾  Transcript + SRT saved", flush=True)

        # ── Per-chunk dual-system analysis ────────────────────────────────────
        chunk_analyses = []   # list of (chunk_start, chunk_dur, ant_result, oai_result)
        for ci, (chunk_start, chunk_dur, chunk_text) in enumerate(chunk_transcripts, 1):
            chunk_end = min(chunk_start + chunk_dur, duration)
            label = f"chunk {ci}/{len(chunk_transcripts)}  ({int(chunk_start//60)}:{int(chunk_start%60):02d}–{int(chunk_end//60)}:{int(chunk_end%60):02d})"

            if not chunk_text.strip():
                print(f"  ⚠  {label}: no speech, skipping", flush=True)
                continue

            step(f"Anthropic  →  {label}")
            t0 = time.time()
            ant_result = analyze_anthropic_chunk(
                ant_client, mov.name, chunk_start, chunk_end, duration, chunk_text
            )
            done("done", time.time() - t0)

            step(f"OpenAI     →  {label}")
            t0 = time.time()
            oai_result = analyze_openai_chunk(
                oai_client, mov.name, chunk_start, chunk_end, duration, chunk_text
            )
            done("done", time.time() - t0)

            chunk_analyses.append((chunk_start, chunk_dur, ant_result, oai_result))

        step(f"Synthesizing {len(chunk_analyses)} chunk analyses → final journal entry")
        t0 = time.time()
        analysis = synthesize_all_chunks(ant_client, mov.name, duration, chunk_analyses)
        done("Synthesis complete", time.time() - t0)

        # ── Extract screenshots (only screenshot_worthy timestamps) ────────────
        ts_set: set = set()
        for km in analysis.get("key_moments", []):
            if km.get("screenshot_worthy", True):
                ts_set.add(float(km.get("timestamp", 0)))
        for topic in analysis.get("topics", []):
            if topic.get("screenshot_worthy", False):
                ts_set.add(float(topic.get("timestamp_start", 0)))

        ts_list = sorted(max(0.0, min(t, duration - 2)) for t in ts_set)
        step(f"Extracting {len(ts_list)} screenshot(s) (visual moments only)")
        t0 = time.time()
        screenshots = []
        for ts in tqdm(ts_list, desc="  screenshots", unit="frame", leave=False):
            out = SCREENSHOTS_DIR / f"{stem}_{int(ts):05d}.jpg"
            if extract_frame(mov, ts, out):
                screenshots.append((ts, out))
        done(f"{len(screenshots)} screenshot(s) saved", time.time() - t0)

        entries.append({
            "video_name":  mov.name,
            "duration":    duration,
            "segments":    segments,
            "analysis":    analysis,
            "screenshots": screenshots,
        })

    if not entries:
        print("\nNo entries with speech — nothing to journal.")
        return

    step("Building journal Word document...")
    t0 = time.time()
    journal_path = JOURNAL_DIR / "video_journal.docx"
    build_journal(entries, journal_path)
    done("Journal saved", time.time() - t0)

    total = time.time() - t0_total
    print(f"\n{'━'*60}")
    print(f"✓  All done!  Total time: {int(total//60)}m {int(total%60)}s")
    print(f"   Subtitles   → {SUBTITLES_DIR}")
    print(f"   Transcripts → {TRANSCRIPTIONS_DIR}")
    print(f"   Journal     → {journal_path}")
    print(f"{'━'*60}")


if __name__ == "__main__":
    main()
